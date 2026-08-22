
# JobTrack
from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from pymongo import MongoClient
from bson import ObjectId
from datetime import datetime, timezone
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
import os
import requests
from urllib.parse import urlparse
import json
import re
from dotenv import load_dotenv
from flask import send_file
import io
load_dotenv()  # loads .env file automatically

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY")
MONGO_URI = os.environ.get("MONGO_URI")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
GEMINI_URL = "https://api.groq.com/openai/v1/chat/completions"
SERPAPI_KEY = os.environ.get("SERPAPI_KEY")


# COMMENT JUST FOR CICD PIPELINING CHECK
#COMMENT JUST FOR CICD PIPELINING CHECK

client = MongoClient(MONGO_URI)
db = client["job_tracker"]
jobs = db["applications"]
users = db["users"]
debriefs = db["debriefs"]
profiles = db["profiles"]
# ── XP Config ────────────────────────────────────────────
XP_RULES = {
    "Applied":   10,
    "Interview": 30,
    "Offer":     100,
    "Rejected":  5,
}


BADGES = [
    {"id": "first_step",     "name": "First Step",        "desc": "Submit your first application",          "icon": "🚀", "condition": lambda s: s["total"] >= 1},
    {"id": "persistence",    "name": "Persistence King",  "desc": "Submit 10 applications",                 "icon": "👑", "condition": lambda s: s["total"] >= 10},
    {"id": "hustler",        "name": "Hustler",           "desc": "Submit 25 applications",                 "icon": "⚡", "condition": lambda s: s["total"] >= 25},
    {"id": "interview_ace",  "name": "Interview Ace",     "desc": "Land your first interview",              "icon": "🎯", "condition": lambda s: s.get("Interview", 0) >= 1},
    {"id": "multi_interview","name": "In Demand",         "desc": "Land 3 interviews",                      "icon": "🔥", "condition": lambda s: s.get("Interview", 0) >= 3},
    {"id": "offer_getter",   "name": "Offer Getter",      "desc": "Receive your first offer",               "icon": "💎", "condition": lambda s: s.get("Offer", 0) >= 1},
    {"id": "resilient",      "name": "Resilient",         "desc": "Keep going after 5 rejections",          "icon": "🛡️", "condition": lambda s: s.get("Rejected", 0) >= 5},
    {"id": "debriefer",      "name": "Debriefer",         "desc": "Complete your first interview debrief",  "icon": "📝", "condition": lambda s: s.get("debriefs", 0) >= 1},
]

def compute_xp_and_badges(uid):
    pipeline = [{"$match": {"user_id": uid}}, {"$group": {"_id": "$status", "count": {"$sum": 1}}}]
    status_counts = {r["_id"]: r["count"] for r in jobs.aggregate(pipeline)}
    total = sum(status_counts.values())
    xp = sum(XP_RULES.get(s, 0) * c for s, c in status_counts.items())
    debrief_count = debriefs.count_documents({"user_id": uid})
    stats = {**status_counts, "total": total, "debriefs": debrief_count}
    earned = [{"id": b["id"], "name": b["name"], "desc": b["desc"], "icon": b["icon"]}
              for b in BADGES if b["condition"](stats)]
    level = max(1, xp // 100)
    xp_in_level = xp % 100
    return {"xp": xp, "level": level, "xp_in_level": xp_in_level, "badges": earned, "status_counts": status_counts, "total": total}

def ghosting_score(job):
    if job.get("status") not in ("Applied", "Interview"):
        return None
    created = job.get("created_at", "")
    try:
        dt = datetime.fromisoformat(created)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        days = (datetime.now(timezone.utc) - dt).days
    except:
        return None
    if days < 7:
        level = "safe"
    elif days < 14:
        level = "amber"
    elif days < 21:
        level = "red"
    else:
        level = "ghost"
    return {"days": days, "level": level}

def gemini(prompt):
    if not GEMINI_API_KEY:
        return "No API key configured."
    payload = {
        "model": "openai/gpt-oss-120b",
        "messages": [{"role": "user", "content": prompt}]
    }
    headers = {
        "Authorization": f"Bearer {GEMINI_API_KEY}",
        "Content-Type": "application/json"
    }
    r = requests.post(GEMINI_URL, json=payload, headers=headers, timeout=60)
    if r.status_code != 200:
        return f"API error: {r.status_code} — {r.text[:200]}"
    try:
        return r.json()["choices"][0]["message"]["content"]
    except:
        return "Could not parse response."
# ── Auth Helpers ──────────────────────────────────────────

def serialize(doc):
    doc["_id"] = str(doc["_id"])
    return doc

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)
    return decorated

def api_login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        return f(*args, **kwargs)
    return decorated

def uid():
    return session.get("user_id")

# ── Pages ─────────────────────────────────────────────────

@app.route("/")
def root():
    if "user_id" in session:
        return redirect(url_for("index"))
    return render_template("landing.html")  # ← change this line

@app.route("/login")
def login_page():
    return redirect(url_for("index")) if "user_id" in session else render_template("login.html")

@app.route("/register")
def register_page():
    return redirect(url_for("index")) if "user_id" in session else render_template("register.html")

@app.route("/dashboard")
@login_required
def index():
    return render_template("index.html", username=session.get("username"))

@app.route("/analytics")
@login_required
def analytics_page():
    return render_template("analytics.html", username=session.get("username"))

@app.route("/debrief/<job_id>")
@login_required
def debrief_page(job_id):
    job = jobs.find_one({"_id": ObjectId(job_id), "user_id": uid()})
    if not job:
        return redirect(url_for("index"))
    existing = debriefs.find_one({"job_id": job_id, "user_id": uid()})
    return render_template("debrief.html", username=session.get("username"),
                           job=serialize(job), existing=serialize(existing) if existing else None)

@app.route("/interview/<job_id>")
@login_required
def interview_page(job_id):
    job = jobs.find_one({"_id": ObjectId(job_id), "user_id": uid()})
    if not job:
        return redirect(url_for("index"))
    return render_template("interview.html", username=session.get("username"), job=serialize(job))

@app.route("/culture/<job_id>")
@login_required
def culture_page(job_id):
    job = jobs.find_one({"_id": ObjectId(job_id), "user_id": uid()})
    if not job:
        return redirect(url_for("index"))
    return render_template("culture.html", username=session.get("username"), job=serialize(job))

# ── Auth API ──────────────────────────────────────────────

@app.route("/api/register", methods=["POST"])
def register():
    data = request.json
    username = data.get("username", "").strip().lower()
    email    = data.get("email", "").strip().lower()
    password = data.get("password", "")
    if not username or not email or not password:
        return jsonify({"error": "All fields required"}), 400
    if len(password) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400
    if users.find_one({"$or": [{"username": username}, {"email": email}]}):
        return jsonify({"error": "Username or email already exists"}), 409
    result = users.insert_one({"username": username, "email": email,
                                "password": generate_password_hash(password),
                                "created_at": datetime.utcnow().isoformat()})
    session["user_id"] = str(result.inserted_id)
    session["username"] = username
    return jsonify({"message": "Registered"}), 201

@app.route("/api/login", methods=["POST"])
def login():
    data = request.json
    identifier = data.get("identifier", "").strip().lower()
    password   = data.get("password", "")
    user = users.find_one({"$or": [{"username": identifier}, {"email": identifier}]})
    if not user or not check_password_hash(user["password"], password):
        return jsonify({"error": "Invalid credentials"}), 401
    session["user_id"] = str(user["_id"])
    session["username"] = user["username"]
    return jsonify({"message": "Logged in", "username": user["username"]})

@app.route("/api/logout", methods=["POST"])
def logout():
    session.clear()
    return jsonify({"message": "Logged out"})

# ── Jobs API ──────────────────────────────────────────────

@app.route("/api/jobs", methods=["GET"])
@api_login_required
def get_jobs():
    all_jobs = list(jobs.find({"user_id": uid()}).sort("created_at", -1))
    result = []
    for j in all_jobs:
        s = serialize(j)
        s["ghosting"] = ghosting_score(j)
        result.append(s)
    return jsonify(result)

@app.route("/api/jobs", methods=["POST"])
@api_login_required
def add_job():
    data = request.json
    job = {
        "user_id":    uid(),
        "company":    data.get("company", "").strip(),
        "role":       data.get("role", "").strip(),
        "status":     data.get("status", "Applied"),
        "location":   data.get("location", "").strip(),
        "salary":     data.get("salary", "").strip(),
        "link":       data.get("link", "").strip(),
        "notes":      data.get("notes", "").strip(),
        "deadline":   data.get("deadline", "").strip(),
        "checklist":  data.get("checklist", []),
        "created_at": datetime.utcnow().isoformat()
    }
    result = jobs.insert_one(job)
    job["_id"] = str(result.inserted_id)
    job["ghosting"] = ghosting_score(job)
    return jsonify(job), 201

@app.route("/api/jobs/<job_id>", methods=["PUT"])
@api_login_required
def update_job(job_id):
    data = request.json
    allowed = ["company", "role", "status", "location", "salary", "link", "notes", "deadline", "checklist"]
    update = {k: data[k] for k in allowed if k in data}
    jobs.update_one({"_id": ObjectId(job_id), "user_id": uid()}, {"$set": update})
    updated = jobs.find_one({"_id": ObjectId(job_id)})
    s = serialize(updated)
    s["ghosting"] = ghosting_score(updated)
    return jsonify(s)

@app.route("/api/jobs/<job_id>", methods=["DELETE"])
@api_login_required
def delete_job(job_id):
    jobs.delete_one({"_id": ObjectId(job_id), "user_id": uid()})
    return jsonify({"deleted": job_id})

# ── Stats & XP ────────────────────────────────────────────

@app.route("/api/stats", methods=["GET"])
@api_login_required
def stats():
    u = uid()
    pipeline = [{"$match": {"user_id": u}}, {"$group": {"_id": "$status", "count": {"$sum": 1}}}]
    status_data = {r["_id"]: r["count"] for r in jobs.aggregate(pipeline)}
    all_jobs = list(jobs.find({"user_id": u}, {"created_at": 1}))
    monthly = {}
    for j in all_jobs:
        try:
            key = datetime.fromisoformat(j["created_at"]).strftime("%b %d")
            monthly[key] = monthly.get(key, 0) + 1
        except: pass
    company_pipe = [{"$match": {"user_id": u}}, {"$group": {"_id": "$company", "count": {"$sum": 1}}}, {"$sort": {"count": -1}}, {"$limit": 5}]
    top_companies = [{"company": r["_id"], "count": r["count"]} for r in jobs.aggregate(company_pipe)]
    total = sum(status_data.values()) or 1
    xp_data = compute_xp_and_badges(u)
    return jsonify({
        "status": status_data,
        "timeline": monthly,
        "top_companies": top_companies,
        "success_rate": round((status_data.get("Offer", 0) / total) * 100, 1),
        "interview_rate": round(((status_data.get("Interview", 0) + status_data.get("Offer", 0)) / total) * 100, 1),
        "total": total,
        "xp": xp_data
    })

@app.route("/api/xp", methods=["GET"])
@api_login_required
def get_xp():
    return jsonify(compute_xp_and_badges(uid()))

# ── Debrief API ───────────────────────────────────────────

@app.route("/api/debrief/<job_id>", methods=["POST"])
@api_login_required
def save_debrief(job_id):
    data = request.json
    doc = {
        "user_id":        uid(),
        "job_id":         job_id,
        "hardest_q":      data.get("hardest_q", ""),
        "excited_about":  data.get("excited_about", ""),
        "went_well":      data.get("went_well", ""),
        "went_poorly":    data.get("went_poorly", ""),
        "followup":       data.get("followup", ""),
        "energy":         data.get("energy", 5),
        "overall":        data.get("overall", 5),
        "saved_at":       datetime.utcnow().isoformat()
    }
    debriefs.update_one({"job_id": job_id, "user_id": uid()}, {"$set": doc}, upsert=True)
    return jsonify({"message": "Debrief saved"})

@app.route("/api/debrief/<job_id>", methods=["GET"])
@api_login_required
def get_debrief(job_id):
    d = debriefs.find_one({"job_id": job_id, "user_id": uid()})
    return jsonify(serialize(d) if d else {})

# ── AI Endpoints ──────────────────────────────────────────

@app.route("/api/ai/culture", methods=["POST"])
@api_login_required
def ai_culture():
    data = request.json
    company  = data.get("company", "")
    role     = data.get("role", "")
    jd       = data.get("job_description", "")
    about    = data.get("about_us", "")
    prefs    = data.get("user_prefs", "")
    prompt = f"""You are a career coach and organizational psychologist. Analyze this company and job posting.

Company: {company}
Role: {role}

Job Description:
{jd}

About Us / Company Culture:
{about}

User's self-reported work preferences:
{prefs}

Provide a structured analysis with these exact sections:

## Culture Vibe
Describe the company culture in 2-3 sentences using vivid adjectives (e.g. "fast-paced and results-driven" or "collaborative and academic").

## Culture Match Score
Rate the match between the user's preferences and this company culture as X/10 with a one-line explanation.

## Language Mirror
List 6-8 specific keywords or phrases the user should naturally weave into their interview answers to resonate with this company. Format as a simple list.

## Watch Out For
2-3 potential culture friction points the user should be aware of.

## Tailored Advice
3 concrete, specific tips for this exact company and role.

Keep the tone honest, direct, and genuinely helpful."""
    return jsonify({"result": gemini(prompt)})

@app.route("/api/ai/interview/start", methods=["POST"])
@api_login_required
def ai_interview_start():
    data = request.json
    company = data.get("company", "")
    role    = data.get("role", "")
    prompt = f"""You are a professional recruiter at {company} interviewing a candidate for the {role} position.

Start the interview naturally. Introduce yourself briefly (make up a name), welcome the candidate, and ask your FIRST interview question. 

Mix behavioral, situational, and role-specific questions throughout the interview.
Ask only ONE question at a time.
Keep your messages concise — like a real interview.
Do not number your questions.
Start now."""
    return jsonify({"result": gemini(prompt)})

@app.route("/api/ai/interview/respond", methods=["POST"])
@api_login_required
def ai_interview_respond():
    data = request.json
    company  = data.get("company", "")
    role     = data.get("role", "")
    history  = data.get("history", [])
    answer   = data.get("answer", "")
    q_count  = data.get("question_count", 1)

    history_text = "\n".join([f"{'Recruiter' if m['role']=='ai' else 'Candidate'}: {m['text']}" for m in history])

    if q_count >= 5:
        prompt = f"""You are a recruiter at {company} interviewing for {role}.

Conversation so far:
{history_text}
Candidate: {answer}

This was the final answer. Wrap up the interview naturally — thank the candidate, mention next steps briefly, and then provide a "Confidence Score" section formatted exactly like this:

---
**Confidence Score: X/10**

**Strengths:** (2-3 bullet points of what came across well)
**Areas to improve:** (2-3 bullet points of constructive feedback)
**Overall impression:** (1-2 sentences)
---"""
    else:
        prompt = f"""You are a recruiter at {company} interviewing for {role}.

Conversation so far:
{history_text}
Candidate: {answer}

React briefly and naturally to their answer (1-2 sentences max), then ask your next interview question. Ask only ONE question. Keep it conversational."""

    return jsonify({"result": gemini(prompt)})


# ── Page Routes ───────────────────────────────────────────────────────────────

@app.route("/coverletter/<job_id>")
@login_required
def coverletter_page(job_id):
    job = jobs.find_one({"_id": ObjectId(job_id), "user_id": uid()})
    if not job:
        return redirect(url_for("index"))
    return render_template("coverletter.html", username=session.get("username"), job=serialize(job))


@app.route("/import")
@login_required
def import_page():
    return render_template("import_job.html", username=session.get("username"))


# ── Cover Letter API ──────────────────────────────────────────────────────────

@app.route("/api/ai/coverletter", methods=["POST"])
@api_login_required
def ai_coverletter():
    data = request.json
    company = data.get("company", "")
    role = data.get("role", "")
    jd = data.get("job_description", "")
    resume = data.get("resume", "")
    name = data.get("name", "The Applicant")
    email = data.get("email", "")
    tone = data.get("tone", "professional")
    length = data.get("length", "medium")
    extra = data.get("extra", "")

    tone_map = {
        "professional": "formal, polished, and confident",
        "enthusiastic": "energetic, passionate, and genuinely excited about the role",
        "concise": "extremely direct and to-the-point — no fluff whatsoever",
        "storytelling": "narrative-driven — open with a brief compelling story or moment",
        "casual": "warm, friendly, and conversational — like writing to a colleague",
    }
    length_map = {
        "short": "Keep it SHORT — around 150 words maximum. One strong paragraph.",
        "medium": "Keep it MEDIUM — around 250 words. Three tight paragraphs.",
        "long": "Write a FULL letter — around 400 words. Four paragraphs with depth.",
    }

    prompt = f"""You are an expert career coach and professional writer. Write a cover letter for this application.

Applicant name: {name}
{f"Applicant email: {email}" if email else ""}
Company: {company}
Role: {role}

Job Description:
{jd if jd else "Not provided — infer requirements from the company name and role."}

Applicant background / resume points:
{resume if resume else "Not provided — write a strong general letter for this role."}

Tone: {tone_map.get(tone, "professional")}
Length: {length_map.get(length, "Around 250 words.")}
{f"Special instructions: {extra}" if extra else ""}

Rules:
- Output ONLY the cover letter text — no meta-commentary, no notes, no explanations
- Do NOT use placeholder brackets like [Your Name] — use the actual name given
- Start with the salutation e.g. "Dear Hiring Manager," or "Dear {company} Team,"
- End with a proper sign-off and the applicant name
- Reference the company name and role specifically — make it feel genuinely tailored
- Avoid the cliche opener "I am writing to express my interest" — be stronger
- No filler phrases like "I am a passionate individual" or "I would be a great fit\""""

    result = gemini(prompt)
    return jsonify({"result": result})


# ── Helpers for import ────────────────────────────────────────────────────────

def _parse_ai_job_json(raw_text):
    """Extract JSON from AI response, handling markdown code fences."""
    import json, re
    text = raw_text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"```\s*$", "", text, flags=re.MULTILINE)
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        return json.loads(match.group())
    raise ValueError("No JSON found in AI response")


# ── Job Import API ────────────────────────────────────────────────────────────

@app.route("/api/ai/import/paste", methods=["POST"])
@api_login_required
def ai_import_paste():
    """Extract job fields from pasted job description text using AI."""
    import re
    data = request.json
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "No text provided"}), 400

    prompt = f"""You are a job posting parser. Extract structured information from this job posting.

Job posting:
{text[:4000]}

Return ONLY a valid JSON object with these exact keys (empty string "" if not found):
{{
  "company": "company name",
  "role": "job title",
  "location": "city, country or Remote",
  "salary": "salary range or empty string",
  "link": "application URL if mentioned, else empty string",
  "notes": "2-3 sentence summary of the role and key requirements"
}}

Return ONLY the JSON. No explanation, no markdown, no code fences."""

    raw = gemini(prompt)
    try:
        extracted = _parse_ai_job_json(raw)
        clean = {k: str(extracted.get(k, "")).strip()
                 for k in ["company", "role", "location", "salary", "link", "notes"]}
        return jsonify(clean)
    except Exception as e:
        return jsonify({"error": f"Could not parse AI response: {str(e)}"}), 500


@app.route("/api/ai/import/url", methods=["POST"])
@api_login_required
def ai_import_url():
    """Fetch a LinkedIn/Indeed URL, extract page text, then parse with AI."""
    import re
    from urllib.parse import urlparse as _urlparse
    data = request.json
    url = data.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400

    try:
        parsed = _urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return jsonify({"error": "Invalid URL"}), 400
    except Exception:
        return jsonify({"error": "Invalid URL"}), 400

    # Fetch page
    try:
        hdrs = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "en-US,en;q=0.9",
        }
        resp = requests.get(url, headers=hdrs, timeout=15)
        resp.raise_for_status()
        page_text = resp.text
    except requests.exceptions.Timeout:
        return jsonify({"error": "Request timed out. Try pasting the description instead."}), 502
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Could not fetch URL: {str(e)[:100]}. Try pasting instead."}), 502

    # Strip HTML
    clean = re.sub(r"<style[^>]*>.*?</style>", " ", page_text, flags=re.DOTALL | re.IGNORECASE)
    clean = re.sub(r"<script[^>]*>.*?</script>", " ", clean, flags=re.DOTALL | re.IGNORECASE)
    clean = re.sub(r"<[^>]+>", " ", clean)
    clean = re.sub(r"\s+", " ", clean).strip()

    if len(clean) < 100:
        return jsonify({
            "error": "Could not extract content. LinkedIn/Indeed may block this. Please paste the description text instead.",
            "company": "", "role": "", "location": "", "salary": "", "link": url, "notes": ""
        }), 200

    prompt = f"""You are a job posting parser. Extract structured information from this webpage text scraped from a job posting.

Webpage text:
{clean[:4000]}

Original URL: {url}

Return ONLY a valid JSON object with these exact keys (empty string "" if not found):
{{
  "company": "company name",
  "role": "job title",
  "location": "city, country or Remote",
  "salary": "salary range or empty string",
  "link": "{url}",
  "notes": "2-3 sentence summary of the role and key requirements"
}}

Return ONLY the JSON. No explanation, no markdown, no code fences."""

    raw = gemini(prompt)
    try:
        extracted = _parse_ai_job_json(raw)
        clean_out = {k: str(extracted.get(k, "")).strip()
                     for k in ["company", "role", "location", "salary", "link", "notes"]}
        if not clean_out.get("link"):
            clean_out["link"] = url
        return jsonify(clean_out)
    except Exception as e:
        return jsonify({"error": f"Could not parse AI response: {str(e)}"}), 500


@app.route("/profile")
@login_required
def profile_page():
    return render_template("profile.html", username=session.get("username"))


# ── Profile API ───────────────────────────────────────────────────────────────

@app.route("/api/profile", methods=["GET"])
@api_login_required
def get_profile():
    """Get the current user's profile."""
    p = profiles.find_one({"user_id": uid()})
    if not p:
        return jsonify({})
    p["_id"] = str(p["_id"])
    return jsonify(p)


@app.route("/api/profile", methods=["POST"])
@api_login_required
def save_profile():
    """Save or update the current user's profile."""
    data = request.json
    doc = {
        "user_id": uid(),
        "personal": data.get("personal", {}),
        "education": data.get("education", []),
        "experience": data.get("experience", []),
        "projects": data.get("projects", []),
        "skills": data.get("skills", []),
        "certifications": data.get("certifications", []),
        "languages": data.get("languages", []),
        "preferences": data.get("preferences", {}),
        "updated_at": datetime.utcnow().isoformat()
    }
    profiles.update_one(
        {"user_id": uid()},
        {"$set": doc},
        upsert=True
    )
    return jsonify({"message": "Profile saved"})


@app.route("/api/profile/summary", methods=["GET"])
@api_login_required
def get_profile_summary():
    """
    Returns a compact profile summary used by CV generator and job matching.
    Calculates completeness percentage.
    """
    p = profiles.find_one({"user_id": uid()})
    if not p:
        return jsonify({"complete": False, "completeness": 0, "missing": ["Full profile"]})

    per = p.get("personal", {})
    prefs = p.get("preferences", {})

    checks = {
        "Name": bool(per.get("name")),
        "Email": bool(per.get("email")),
        "Summary": bool(per.get("summary")),
        "Education": len(p.get("education", [])) > 0,
        "Skills": len(p.get("skills", [])) > 0,
        "Projects": len(p.get("projects", [])) > 0,
        "Job Type": bool(prefs.get("job_type")),
        "Location": bool(prefs.get("location")),
    }

    done = sum(1 for v in checks.values() if v)
    pct = round((done / len(checks)) * 100)
    missing = [k for k, v in checks.items() if not v]

    return jsonify({
        "complete": pct >= 70,
        "completeness": pct,
        "missing": missing,
        "name": per.get("name", ""),
        "skills": p.get("skills", []),
        "preferences": prefs,
    })


# ── Job Feed Page Route ───────────────────────────────────────────────────────
@app.route("/jobs/feed")
@login_required
def job_feed_page():
    return render_template("feed.html", username=session.get("username"))


@app.route("/api/jobs/feed", methods=["POST"])
@api_login_required
def fetch_job_feed():
    data            = request.json or {}
    query           = data.get("query", "").strip()
    location        = data.get("location", "").strip()
    employment_type = data.get("employment_type", "")
    date_posted     = data.get("date_posted", "week")

    # Get user profile
    user_profile = profiles.find_one({"user_id": uid()})

    # Build query from profile if nothing given
    if not query and user_profile:
        prefs  = user_profile.get("preferences", {})
        roles  = prefs.get("roles", "")
        skills = user_profile.get("skills", [])
        query  = roles.split(",")[0].strip() if roles else (skills[0] if skills else "Software Engineer")
        if not location:
            location = prefs.get("location", "Pakistan")
        if not employment_type:
            type_map = {"fulltime": "full_time", "parttime": "part_time",
                        "internship": "intern", "contract": "contractor"}
            employment_type = type_map.get(prefs.get("job_type", ""), "")

    if not query:
        query = "Software Engineer"

    # Map date filter
    chips_map = {
        "today":  "date_posted:today",
        "3days":  "date_posted:3days",
        "week":   "date_posted:week",
        "month":  "date_posted:month",
    }
    date_chip = chips_map.get(date_posted, "date_posted:week")

    # Add employment type chip
    type_chip = ""
    if employment_type == "full_time":
        type_chip = "employment_type:FULLTIME"
    elif employment_type == "part_time":
        type_chip = "employment_type:PARTTIME"
    elif employment_type == "intern":
        type_chip = "employment_type:INTERN"
    elif employment_type == "contractor":
        type_chip = "employment_type:CONTRACTOR"

    chips = date_chip
    if type_chip:
        chips += f",{type_chip}"

    # Build SerpAPI params
    params = {
        "engine":   "google_jobs",
        "q":        f"{query} {location}".strip(),
        "hl":       "en",
        "chips":    chips,
        "api_key":  SERPAPI_KEY,
        "num":      20,
    }

    try:
        resp = requests.get(
            "https://serpapi.com/search",
            params=params,
            timeout=20
        )
        resp.raise_for_status()
        result_data = resp.json()
        raw_jobs = result_data.get("jobs_results", [])
    except requests.exceptions.Timeout:
        return jsonify({"error": "Search timed out. Please try again."}), 502
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to fetch jobs: {str(e)[:150]}"}), 502
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {str(e)[:100]}"}), 500

    if not raw_jobs:
        return jsonify({"jobs": [], "tracked_ids": []})

    # Normalize SerpAPI Google Jobs response
    normalized = []
    for j in raw_jobs[:20]:
        # Get salary from extensions
        extensions = j.get("detected_extensions", {})
        salary = extensions.get("salary", "")

        # Get posted date
        posted_str = extensions.get("posted_at", "")
        if not posted_str:
            posted_str = j.get("job_highlights", [{}])[0].get("items", [""])[0] if j.get("job_highlights") else ""

        # Get description from highlights
        desc = ""
        for highlight in j.get("job_highlights", []):
            items = highlight.get("items", [])
            if items:
                desc = " | ".join(items[:3])
                break
        if not desc:
            desc = j.get("description", "")[:500]

        # Get apply link
        apply_link = ""
        apply_options = j.get("apply_options", [])
        if apply_options:
            apply_link = apply_options[0].get("link", "")

        normalized.append({
            "job_id":          j.get("job_id", str(len(normalized))),
            "title":           j.get("title", ""),
            "company":         j.get("company_name", ""),
            "location":        j.get("location", ""),
            "employment_type": extensions.get("work_from_home", False) and "Remote" or
                               extensions.get("schedule_type", ""),
            "salary":          salary,
            "description":     desc[:500],
            "url":             apply_link,
            "posted":          posted_str,
            "posted_at":       posted_str,
            "match_score":     None,
            "match_reason":    None,
        })

    # AI Match Scoring
    if user_profile and user_profile.get("skills"):
        normalized = _score_jobs_with_ai(normalized, user_profile)

    # Check already tracked
    existing_jobs = list(jobs.find({"user_id": uid()}, {"link": 1}))
    tracked_links = {j.get("link", "") for j in existing_jobs}
    tracked_ids   = [j["job_id"] for j in normalized if j["url"] and j["url"] in tracked_links]

    return jsonify({"jobs": normalized, "tracked_ids": tracked_ids})
def _score_jobs_with_ai(job_list, user_profile):
    """Score all jobs against user profile in one AI call."""
    skills      = user_profile.get("skills", [])
    experience  = user_profile.get("experience", [])
    projects    = user_profile.get("projects", [])
    prefs       = user_profile.get("preferences", {})
    education   = user_profile.get("education", [])

    # Build compact profile string
    profile_str = f"""Skills: {', '.join(skills[:20])}
Experience: {'; '.join([f"{e.get('role','')} at {e.get('company','')}" for e in experience[:3]])}
Projects: {'; '.join([f"{p.get('name','')} ({p.get('tech','')})" for p in projects[:3]])}
Education: {'; '.join([f"{e.get('degree','')} at {e.get('school','')}" for e in education[:2]])}
Preferred job type: {prefs.get('job_type', 'any')}
Preferred location: {prefs.get('location', 'any')}
Target roles: {prefs.get('roles', 'any')}"""

    # Build jobs list string
    jobs_str = "\n".join([
        f"{i+1}. {j['title']} at {j['company']} ({j['location']}) — {j['description'][:200]}"
        for i, j in enumerate(job_list)
    ])

    prompt = f"""You are a career matching AI. Score how well each job matches this candidate's profile.

CANDIDATE PROFILE:
{profile_str}

JOBS TO SCORE:
{jobs_str}

Return ONLY a valid JSON array with exactly {len(job_list)} objects, one per job, in the same order:
[
  {{"score": 85, "reason": "Strong Python/Flask match, remote aligns with preference"}},
  {{"score": 60, "reason": "Partial match — requires React which candidate lacks"}},
  ...
]

Rules:
- score is 0-100 integer
- reason is max 10 words explaining the match
- Return ONLY the JSON array, no explanation, no markdown"""

    try:
        raw = gemini(prompt)
        # Clean response
        raw = raw.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE)
        # Find JSON array
        match = re.search(r"\[.*\]", raw, re.DOTALL)
        if not match:
            return job_list
        scores = json.loads(match.group())
        for i, job in enumerate(job_list):
            if i < len(scores):
                job["match_score"]  = int(scores[i].get("score", 0))
                job["match_reason"] = str(scores[i].get("reason", ""))
    except Exception:
        pass  # return unscored jobs if AI fails

    return job_list





def _format_posted(iso_str):
    if not iso_str:
        return ""
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        days = (datetime.now(timezone.utc) - dt).days
        if days == 0:
            return "Today"
        if days == 1:
            return "Yesterday"
        if days < 7:
            return f"{days} days ago"
        if days < 30:
            return f"{days // 7}w ago"
        return dt.strftime("%b %d")
    except Exception:
        return ""




# ── CV Generator Page Route ───────────────────────────────────────────────────

@app.route("/cv/<job_id>")
@login_required
def cv_page(job_id):
    job = jobs.find_one({"_id": ObjectId(job_id), "user_id": uid()})
    if not job:
        return redirect(url_for("index"))
    return render_template("cv.html", username=session.get("username"), job=serialize(job))


# ── CV Generator API ──────────────────────────────────────────────────────────

@app.route("/api/ai/cv/generate", methods=["POST"])
@api_login_required
def ai_cv_generate():
    """
    1. Load user profile from MongoDB
    2. Get job description from request
    3. AI tailors the profile to match the JD (ATS keywords)
    4. Returns structured CV data for preview + download
    """
    data     = request.json or {}
    jd       = data.get("job_description", "").strip()
    length   = data.get("length", "1.5")

    # Load full profile
    profile = profiles.find_one({"user_id": uid()})
    if not profile:
        return jsonify({"error": "Please complete your profile first before generating a CV."}), 400

    per    = profile.get("personal", {})
    exp    = profile.get("experience", [])
    proj   = profile.get("projects", [])
    edu    = profile.get("education", [])
    skills = profile.get("skills", [])
    certs  = profile.get("certifications", [])
    langs  = profile.get("languages", [])

    if not per.get("name"):
        return jsonify({"error": "Please add your name to your profile first."}), 400

    # Build profile summary for AI
    profile_text = f"""
PERSONAL:
Name: {per.get('name', '')}
Email: {per.get('email', '')}
Phone: {per.get('phone', '')}
Location: {per.get('location', '')}
LinkedIn: {per.get('linkedin', '')}
GitHub: {per.get('github', '')}
Summary: {per.get('summary', '')}

EDUCATION:
{chr(10).join([f"- {e.get('degree','')} at {e.get('school','')} ({e.get('start','')}–{e.get('end','')}) GPA: {e.get('gpa','')}" for e in edu])}

EXPERIENCE:
{chr(10).join([f"- {e.get('role','')} at {e.get('company','')} ({e.get('start','')}–{e.get('end','Present')}): {e.get('description','')}" for e in exp])}

PROJECTS:
{chr(10).join([f"- {p.get('name','')} ({p.get('tech','')}): {p.get('description','')}" for p in proj])}

SKILLS: {', '.join(skills)}
CERTIFICATIONS: {', '.join(certs)}
LANGUAGES: {', '.join(langs)}
"""

    length_instruction = "Keep the CV to exactly 1 page — be concise." if length == "1" else \
                         "Keep the CV to 1 to 1.5 pages — comprehensive but tight."

    prompt = f"""You are an expert ATS-optimized CV writer. Create a tailored CV for this candidate.

CANDIDATE PROFILE:
{profile_text}

JOB DESCRIPTION TO OPTIMIZE FOR:
{jd[:3000] if jd else "General professional CV — use best judgment for the candidate's field."}

TASK:
1. Rewrite the professional summary to directly address this role
2. Reorder and rewrite experience bullet points to highlight most relevant achievements using strong action verbs and metrics
3. Reorder projects to show most relevant first
4. Organize skills by category (Languages, Frameworks, Tools, Cloud/DevOps, etc.)
5. Extract 8-12 ATS keywords from the job description to naturally include
6. {length_instruction}

Return ONLY a valid JSON object with this exact structure:
{{
  "personal": {{
    "name": "full name",
    "tagline": "one line role/title matching the job",
    "email": "email",
    "phone": "phone",
    "location": "city, country",
    "linkedin": "linkedin url or username",
    "github": "github url or username"
  }},
  "summary": "2-3 sentence tailored professional summary",
  "skills": {{
    "Languages": ["Python", "JavaScript"],
    "Frameworks": ["Flask", "React"],
    "Tools": ["Docker", "Git"],
    "Cloud & DevOps": ["Azure", "Kubernetes", "CI/CD"]
  }},
  "experience": [
    {{
      "role": "Job Title",
      "company": "Company Name",
      "location": "City",
      "start": "Jan 2024",
      "end": "Present",
      "bullets": [
        "Strong action verb + what you did + measurable result",
        "Another achievement"
      ]
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "tech": "Flask, MongoDB, Docker",
      "url": "github.com/...",
      "year": "2024",
      "bullets": [
        "What it does and impact",
        "Key technical achievement"
      ]
    }}
  ],
  "education": [
    {{
      "degree": "BS Computer Science",
      "school": "COMSATS University",
      "location": "Lahore",
      "start": "2023",
      "end": "2027",
      "gpa": "3.5"
    }}
  ],
  "certifications": ["cert1", "cert2"],
  "languages": ["English (Fluent)", "Urdu (Native)"],
  "ats_keywords": ["keyword1", "keyword2", "keyword3"]
}}

Rules:
- Return ONLY the JSON, no markdown, no explanation
- Every bullet point must start with a strong action verb (Built, Designed, Deployed, Implemented, etc.)
- Include metrics where possible (reduced by X%, served X users, etc.)
- Keep bullets concise — 1 line each
- Skills must be organized into exactly the categories that make sense for this candidate
- ATS keywords must come directly from the job description"""

    raw = gemini(prompt)

    # Parse JSON
    try:
        raw = raw.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"```\s*$", "", raw, flags=re.MULTILINE)
        match_obj = re.search(r"\{.*\}", raw, re.DOTALL)
        if not match_obj:
            return jsonify({"error": "AI could not generate CV structure. Please try again."}), 500
        cv_data = json.loads(match_obj.group())
    except Exception as e:
        return jsonify({"error": f"Failed to parse CV data: {str(e)}"}), 500

    return jsonify({"cv_data": cv_data})


# ── CV Download (PDF or DOCX) ─────────────────────────────────────────────────

@app.route("/api/ai/cv/download", methods=["POST"])
@api_login_required
def ai_cv_download():
    """Generate and return the CV as PDF or DOCX file."""
    data     = request.json or {}
    cv_data  = data.get("cv_data", {})
    fmt      = data.get("format", "pdf")
    filename = data.get("filename", "CV")

    if not cv_data:
        return jsonify({"error": "No CV data provided"}), 400

    if fmt == "pdf":
        pdf_buffer = _generate_cv_pdf(cv_data)
        return send_file(
            pdf_buffer,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{filename}.pdf"
        )
    else:
        docx_buffer = _generate_cv_docx(cv_data)
        return send_file(
            docx_buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{filename}.docx"
        )


def _generate_cv_pdf(cv):
    """Generate ATS-compatible PDF CV using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buffer = io.BytesIO()

    # Page setup — tight margins for 1-1.5 page
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=15*mm,
        rightMargin=15*mm,
        topMargin=14*mm,
        bottomMargin=14*mm,
    )

    # Colors
    BLACK   = HexColor('#1a1a1a')
    DARK    = HexColor('#333333')
    MEDIUM  = HexColor('#555555')
    LIGHT   = HexColor('#777777')
    ACCENT  = HexColor('#1a1a1a')

    # Styles
    def style(name, **kwargs):
        return ParagraphStyle(name, **kwargs)

    s_name = style('name', fontName='Helvetica-Bold', fontSize=18, textColor=BLACK, spaceAfter=1*mm)
    s_tagline = style('tagline', fontName='Helvetica', fontSize=10, textColor=MEDIUM, spaceAfter=2*mm)
    s_contact = style('contact', fontName='Helvetica', fontSize=8.5, textColor=DARK, spaceAfter=3*mm)
    s_sec_title = style('sec_title', fontName='Helvetica-Bold', fontSize=9, textColor=BLACK,
                        spaceBefore=4*mm, spaceAfter=1.5*mm, textTransform='uppercase', letterSpacing=0.5)
    s_summary = style('summary', fontName='Helvetica', fontSize=9.5, textColor=DARK,
                      leading=14, spaceAfter=2*mm)
    s_item_title = style('item_title', fontName='Helvetica-Bold', fontSize=10, textColor=BLACK, spaceAfter=0.5*mm)
    s_item_sub = style('item_sub', fontName='Helvetica-Oblique', fontSize=9, textColor=MEDIUM, spaceAfter=0.5*mm)
    s_bullet = style('bullet', fontName='Helvetica', fontSize=9, textColor=DARK,
                     leading=13, leftIndent=8*mm, firstLineIndent=-4*mm, spaceAfter=1*mm)
    s_skill_cat = style('skill_cat', fontName='Helvetica-Bold', fontSize=9, textColor=BLACK, spaceAfter=1*mm)
    s_skill_val = style('skill_val', fontName='Helvetica', fontSize=9, textColor=DARK, spaceAfter=2*mm)
    s_date = style('date', fontName='Helvetica', fontSize=9, textColor=LIGHT, alignment=2)

    story = []
    per = cv.get('personal', {})

    # Name
    story.append(Paragraph(per.get('name', ''), s_name))

    # Tagline
    if per.get('tagline'):
        story.append(Paragraph(per['tagline'], s_tagline))

    # Contact line
    contact_parts = []
    if per.get('email'):    contact_parts.append(per['email'])
    if per.get('phone'):    contact_parts.append(per['phone'])
    if per.get('location'): contact_parts.append(per['location'])
    if per.get('linkedin'): contact_parts.append(per['linkedin'])
    if per.get('github'):   contact_parts.append(per['github'])
    if contact_parts:
        story.append(Paragraph('  |  '.join(contact_parts), s_contact))

    story.append(HRFlowable(width="100%", thickness=1.5, color=BLACK, spaceAfter=3*mm))

    # Summary
    if cv.get('summary'):
        story.append(Paragraph('PROFESSIONAL SUMMARY', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        story.append(Paragraph(cv['summary'], s_summary))

    # Skills
    skills = cv.get('skills', {})
    if skills:
        story.append(Paragraph('TECHNICAL SKILLS', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        for cat, items in skills.items():
            if items:
                story.append(Paragraph(f"<b>{cat}:</b> {', '.join(items)}", s_skill_val))

    # Experience
    exp = cv.get('experience', [])
    if exp:
        story.append(Paragraph('WORK EXPERIENCE', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        for e in exp:
            date_str = f"{e.get('start', '')}{'  –  ' + e.get('end', '') if e.get('end') else ''}"
            story.append(Paragraph(
                f"<b>{e.get('role', '')}</b>  |  {e.get('company', '')}"
                f"{'  ·  ' + e.get('location', '') if e.get('location') else ''}"
                f"  <font color='#888888' size='8'>{date_str}</font>",
                s_item_title
            ))
            for b in e.get('bullets', []):
                story.append(Paragraph(f"• {b}", s_bullet))
            story.append(Spacer(1, 1*mm))

    # Projects
    proj = cv.get('projects', [])
    if proj:
        story.append(Paragraph('PROJECTS', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        for p in proj:
            tech_str = f"  |  <i>{p.get('tech', '')}</i>" if p.get('tech') else ''
            year_str = f"  <font color='#888888' size='8'>{p.get('year', '')}</font>" if p.get('year') else ''
            story.append(Paragraph(
                f"<b>{p.get('name', '')}</b>{tech_str}{year_str}",
                s_item_title
            ))
            if p.get('url'):
                story.append(Paragraph(p['url'], s_item_sub))
            for b in p.get('bullets', []):
                story.append(Paragraph(f"• {b}", s_bullet))
            story.append(Spacer(1, 1*mm))

    # Education
    edu = cv.get('education', [])
    if edu:
        story.append(Paragraph('EDUCATION', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        for e in edu:
            date_str = f"{e.get('start', '')}{'  –  ' + e.get('end', '') if e.get('end') else ''}"
            gpa_str = f"  ·  GPA: {e['gpa']}" if e.get('gpa') else ''
            story.append(Paragraph(
                f"<b>{e.get('degree', '')}</b>  |  {e.get('school', '')}"
                f"{'  ·  ' + e.get('location', '') if e.get('location') else ''}{gpa_str}"
                f"  <font color='#888888' size='8'>{date_str}</font>",
                s_item_title
            ))
            story.append(Spacer(1, 1.5*mm))

    # Certifications
    certs = cv.get('certifications', [])
    if certs:
        story.append(Paragraph('CERTIFICATIONS', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        story.append(Paragraph('  |  '.join(certs), s_skill_val))

    # Languages
    langs = cv.get('languages', [])
    if langs:
        story.append(Paragraph('LANGUAGES', s_sec_title))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#dddddd'), spaceAfter=2*mm))
        story.append(Paragraph('  |  '.join(langs), s_skill_val))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _generate_cv_docx(cv):
    """Generate ATS-compatible Word CV using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # Fall back to PDF if docx not installed
        return _generate_cv_pdf(cv)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    per = cv.get('personal', {})

    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(per.get('name', ''))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tagline
    if per.get('tagline'):
        tl = doc.add_paragraph()
        tl_run = tl.add_run(per['tagline'])
        tl_run.font.size = Pt(10)
        tl_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Contact
    contact_parts = [per.get(k, '') for k in ['email','phone','location','linkedin','github'] if per.get(k)]
    if contact_parts:
        cp = doc.add_paragraph()
        cr = cp.add_run('  |  '.join(contact_parts))
        cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def add_section(title, content_fn):
        # Section title
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        p.paragraph_format.space_before = Pt(8)
        # HR line via border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
        content_fn()

    # Summary
    if cv.get('summary'):
        def add_summary():
            p = doc.add_paragraph(cv['summary'])
            p.runs[0].font.size = Pt(9.5)
        add_section('PROFESSIONAL SUMMARY', add_summary)

    # Skills
    skills = cv.get('skills', {})
    if skills:
        def add_skills():
            for cat, items in skills.items():
                if items:
                    p = doc.add_paragraph()
                    r1 = p.add_run(f"{cat}: ")
                    r1.bold = True
                    r1.font.size = Pt(9)
                    r2 = p.add_run(', '.join(items))
                    r2.font.size = Pt(9)
        add_section('TECHNICAL SKILLS', add_skills)

    # Experience
    exp = cv.get('experience', [])
    if exp:
        def add_exp():
            for e in exp:
                p = doc.add_paragraph()
                r = p.add_run(f"{e.get('role','')}  |  {e.get('company','')}")
                r.bold = True
                r.font.size = Pt(10)
                date_r = p.add_run(f"  {e.get('start','')}–{e.get('end','')}")
                date_r.font.size = Pt(8.5)
                date_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                for b in e.get('bullets', []):
                    bp = doc.add_paragraph(b, style='List Bullet')
                    bp.runs[0].font.size = Pt(9)
        add_section('WORK EXPERIENCE', add_exp)

    # Projects
    proj = cv.get('projects', [])
    if proj:
        def add_proj():
            for p_item in proj:
                p = doc.add_paragraph()
                r = p.add_run(p_item.get('name', ''))
                r.bold = True
                r.font.size = Pt(10)
                if p_item.get('tech'):
                    tr = p.add_run(f"  |  {p_item['tech']}")
                    tr.font.size = Pt(9)
                    tr.italic = True
                for b in p_item.get('bullets', []):
                    bp = doc.add_paragraph(b, style='List Bullet')
                    bp.runs[0].font.size = Pt(9)
        add_section('PROJECTS', add_proj)

    # Education
    edu = cv.get('education', [])
    if edu:
        def add_edu():
            for e in edu:
                p = doc.add_paragraph()
                r = p.add_run(f"{e.get('degree','')}  |  {e.get('school','')}")
                r.bold = True
                r.font.size = Pt(10)
                date_r = p.add_run(f"  {e.get('start','')}–{e.get('end','')}")
                date_r.font.size = Pt(8.5)
                date_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        add_section('EDUCATION', add_edu)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
